from docx import Document
import PyPDF2

class TextUtils:
    def extract_text_from_docx(filepath):
        """Extract text from a DOCX file."""
        # document = Document(filepath)
        # text = '\n'.join([paragraph.text for paragraph in document.paragraphs])
        # return text
        document = Document(filepath)
        text = '\n'.join([paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()])
        return text.strip()

    def extract_text_from_pdf(filepath):
        """Extract text from a PDF file."""
        text = ''
        with open(filepath, 'rb') as pdf_file:
            reader = PyPDF2.PdfReader(pdf_file)
            for page in reader.pages:
                text += page.extract_text() + '\n'
        return text